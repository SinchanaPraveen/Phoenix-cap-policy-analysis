"""
exporter.py – Export a full pipeline result to .docx or .md format.
"""
from __future__ import annotations

import os
import tempfile
import uuid
from datetime import date
from pathlib import Path
from typing import Any

# Outputs directory sits next to this file (backend/outputs/)
_OUTPUTS_DIR = Path(__file__).resolve().parent / "outputs"


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _outputs_dir() -> Path:
    """Return a writable outputs directory for the current runtime."""
    if os.getenv("VERCEL"):
        # Vercel functions cannot write into the project directory.
        out_dir = Path(tempfile.gettempdir()) / "phoenix-cap-outputs"
    else:
        out_dir = _OUTPUTS_DIR
    out_dir.mkdir(parents=True, exist_ok=True)
    return out_dir


def _run_id(result: dict[str, Any]) -> str:
    """
    Derive a stable run-id from the result dict.

    Uses ``result["run_id"]`` when present; otherwise generates a short UUID.
    """
    return result.get("run_id") or uuid.uuid4().hex[:12]


# ---------------------------------------------------------------------------
# Markdown export
# ---------------------------------------------------------------------------

def export_markdown(result: dict[str, Any]) -> str:
    """
    Build and return a complete Markdown document for *result*.

    Covers all stages present in the result (MVP or Full mode).

    Parameters
    ----------
    result:
        The pipeline result dict returned by :func:`orchestrator.run_pipeline`.

    Returns
    -------
    str
        Formatted Markdown string.
    """
    mode: str = result.get("mode", "mvp")
    idea: str = result.get("stage1", {}).get("idea", "N/A")
    today = date.today().isoformat()
    memo: dict = result.get("memo") or result.get("stage4") or {}
    stage1: dict = result.get("stage1", {})
    stage2: dict = result.get("stage2", {})

    lines: list[str] = [
        "# Phoenix CAP Policy Analysis",
        "",
        f"**Policy Idea:** {idea}  ",
        f"**Date:** {today}  ",
        f"**Mode:** {mode.upper()}  ",
        "",
        "---",
        "",
    ]

    # ── Policy Memo ──────────────────────────────────────────────────────────
    lines += [
        "## POLICY MEMO",
        "",
        f"**TO:** {memo.get('to_recipient', '')}  ",
        f"**CC:** {memo.get('cc_recipient', '')}  ",
        f"**DATE:** {memo.get('date', today)}  ",
        f"**RE:** {memo.get('re_subject', '')}  ",
        "",
        "### Issue",
        memo.get("issue", ""),
        "",
        "### Background",
        memo.get("background", ""),
        "",
        "### Key Findings",
    ]
    for finding in memo.get("key_findings", []):
        lines.append(f"- {finding}")
    lines += [
        "",
        "### Recommendation",
        memo.get("recommendation", ""),
        "",
        "### Next Steps",
    ]
    for i, step in enumerate(memo.get("next_steps", []), start=1):
        lines.append(f"{i}. {step}")
    if memo.get("citations"):
        lines += ["", "### Citations"]
        for c in memo["citations"]:
            lines.append(f"- {c}")
    lines += [
        "",
        f"*Word count: {memo.get('word_count', 0)}*",
        "",
        "---",
        "",
    ]

    # ── Stage 1 Assessment ───────────────────────────────────────────────────
    lines += [
        "## STAGE 1 ASSESSMENT",
        "",
        stage1.get("raw_output", ""),
        "",
        "---",
        "",
    ]

    # ── Stage 2 Roadmap ──────────────────────────────────────────────────────
    lines += [
        "## STAGE 2 ROADMAP",
        "",
        stage2.get("raw_output", ""),
        "",
        "---",
        "",
    ]

    # ── Full-mode sections ───────────────────────────────────────────────────
    if mode == "full":
        for key, title in [
            ("stage3a", "STAGE 3A – NEUTRAL EVALUATOR"),
            ("stage3b", "STAGE 3B – CITIZEN EVALUATOR"),
            ("stage3c", "STAGE 3C – ADEQ EVALUATOR"),
        ]:
            stage_data: dict = result.get(key, {})
            lines += [
                f"## {title}",
                "",
                f"**Agent:** {stage_data.get('agent_name', '')}",
                "",
                stage_data.get("raw_output", ""),
                "",
                "---",
                "",
            ]

        synthesis: dict = result.get("synthesis", {})
        lines += [
            "## STAGE 5 – SYNTHESIS",
            "",
            synthesis.get("raw_output", ""),
            "",
            "---",
            "",
        ]

    # ── Weights ──────────────────────────────────────────────────────────────
    weights: dict = result.get("stage1", {})  # weights echoed from IdeaInput
    # Weights are not stored in stage output; skip silently if absent.

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------

def export_docx(result: dict[str, Any]) -> str:
    """
    Build a formatted .docx document for *result* and save it to disk.

    Parameters
    ----------
    result:
        The pipeline result dict returned by :func:`orchestrator.run_pipeline`.

    Returns
    -------
    str
        Absolute path to the saved ``.docx`` file.

    Raises
    ------
    ImportError
        If ``python-docx`` is not installed.
    """
    try:
        from docx import Document  # noqa: PLC0415
        from docx.shared import Pt  # noqa: PLC0415
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: PLC0415
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for .docx export. "
            "Run: pip install python-docx"
        ) from exc

    mode: str = result.get("mode", "mvp")
    today = date.today().isoformat()
    memo: dict = result.get("memo") or result.get("stage4") or {}
    stage1: dict = result.get("stage1", {})
    stage2: dict = result.get("stage2", {})
    idea: str = stage1.get("idea", "N/A")

    doc = Document()

    # ── Title block ──────────────────────────────────────────────────────────
    title_para = doc.add_heading("Phoenix CAP Policy Analysis", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f"{idea}\n{today}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    doc.add_paragraph()  # spacer

    # ── Helper closures ───────────────────────────────────────────────────────
    def _h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def _p(text: str) -> None:
        doc.add_paragraph(text or "")

    def _bullet(items: list[str]) -> None:
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    def _numbered(items: list[str]) -> None:
        for item in items:
            doc.add_paragraph(item, style="List Number")

    def _bold_field(label: str, value: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(f"{label}: ")
        run.bold = True
        para.add_run(value or "")

    # ── Section: Policy Memo ─────────────────────────────────────────────────
    _h("POLICY MEMO")
    _bold_field("TO", memo.get("to_recipient", ""))
    _bold_field("CC", memo.get("cc_recipient", ""))
    _bold_field("DATE", memo.get("date", today))
    _bold_field("RE", memo.get("re_subject", ""))
    doc.add_paragraph()

    _h("Issue", level=2)
    _p(memo.get("issue", ""))

    _h("Background", level=2)
    _p(memo.get("background", ""))

    _h("Key Findings", level=2)
    _bullet(memo.get("key_findings", []))

    _h("Recommendation", level=2)
    _p(memo.get("recommendation", ""))

    _h("Next Steps", level=2)
    _numbered(memo.get("next_steps", []))

    if memo.get("citations"):
        _h("Citations", level=2)
        _bullet(memo["citations"])

    _p(f"Word count: {memo.get('word_count', 0)}")

    # ── Full-mode sections ───────────────────────────────────────────────────
    if mode == "full":
        for key, title in [
            ("stage3a", "STAGE 3A – NEUTRAL EVALUATOR"),
            ("stage3b", "STAGE 3B – CITIZEN EVALUATOR"),
            ("stage3c", "STAGE 3C – ADEQ EVALUATOR"),
        ]:
            stage_data: dict = result.get(key, {})
            doc.add_page_break()
            _h(title)
            agent_name = stage_data.get("agent_name", "")
            if agent_name:
                _p(f"Agent: {agent_name}")
            _p(stage_data.get("raw_output", ""))

        synthesis: dict = result.get("synthesis", {})
        doc.add_page_break()
        _h("STAGE 5 – SYNTHESIS")
        _p(synthesis.get("raw_output", ""))

    # ── Section: Evaluation Weights ──────────────────────────────────────────
    weights_raw: dict = result.get("weights", {})
    if weights_raw:
        doc.add_page_break()
        _h("EVALUATION WEIGHTS USED")
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Criterion"
        hdr[1].text = "Weight (0–10)"
        for hdr_cell in hdr:
            for run in hdr_cell.paragraphs[0].runs:
                run.bold = True
        criteria_labels = {
            "effectiveness": "Effectiveness",
            "efficiency": "Efficiency",
            "equity": "Equity",
            "liberty": "Liberty",
            "political_feasibility": "Political Feasibility",
            "social_acceptability": "Social Acceptability",
            "administrative_feasibility": "Administrative Feasibility",
        }
        for key, label in criteria_labels.items():
            value = weights_raw.get(key)
            if value is not None:
                row = table.add_row().cells
                row[0].text = label
                row[1].text = str(value)

    # ── Save ─────────────────────────────────────────────────────────────────
    run_id = _run_id(result)
    out_path = _outputs_dir() / f"{run_id}.docx"
    doc.save(str(out_path))
    return str(out_path)
